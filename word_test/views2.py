from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WD_ALIGN_PARAGRAPH   # 用作设置段落对齐
from docx.shared import Pt  # 磅数
from docx.oxml.ns import qn  # 中文格式
from docx.shared import Inches  # 图片尺寸

import time
today = time.strftime("%Y-%m-%d",time.localtime())
price = 100

if __name__ == '__main__':
    document = Document()
    document.styles['Normal'].font.name=u'微软雅黑'
    document.styles['Normal'].font.size=Pt(14)
    # 设置文档的基础字体
    document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

    p1 = document.add_paragraph()
    # 初始化建立第一个自然段  设置对齐方式为居中，默认为左对齐
    p1.aligment = WD_ALIGN_PARAGRAPH.CENTER
    # 标题内容
    run1 = p1.add_run('人像细节特征比较图：(主要五官的特征比对)')
    run1.font.name = '微软雅黑'
    run1.font.size = Pt(21)  # 设置字体
    # 设置加粗
    run1.font.bold = True
    # 段后距离5磅
    p1.space_after = Pt(5)
    ##段前距离5磅
    p1.space_before = Pt(5)

    table = document.add_table(rows=5, cols=4, style='Table Grid')

    table.cell(0, 0).text = '名称'
    table.cell(0, 1).text = '检材细节特征'
    table.cell(0, 2).text = '样本细节特征'
    table.cell(0, 3).text = '细节特征一致性对比结果'
    table.cell(1, 0).text = '眼睛及眉毛特征'
    table.cell(2, 0).text = '鼻部及人中特征'
    table.cell(3, 0).text = '唇部特征'
    table.cell(4, 0).text = '耳部特征'
    run = document.tables[0].cell(1, 2).paragraphs[0].add_run()
    run.add_picture('2.png')

    document.save('测试文档2.docx')  # 以“客户名-价格通知”作为文件名保存



