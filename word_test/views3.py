from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WD_ALIGN_PARAGRAPH   # 用作设置段落对齐
from docx.shared import Pt, RGBColor  # 磅数
from docx.oxml.ns import qn  # 中文格式
from docx.shared import Inches  # 图片尺寸


import time
today=time.strftime("%Y-%m-%d",time.localtime())


if __name__ == '__main__':
    document =Document()
    document.styles['Normal'].font.name = u'微软雅黑'
    document.styles['Normal'].font.size = Pt(14)
    # 设置文档的基础字体
    document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'),u'微软雅黑')



    p1 = document.add_paragraph()
    # 初始化建立第一个自然段  设置对齐方式为居中，默认为左对齐
    p1.aligment=WD_ALIGN_PARAGRAPH.CENTER
    # 标题内容
    run1=p1.add_run('样本中人像头部和五官测量图')
    run1.font.name = '微软雅黑'
    run1.font.size=Pt(21) #设置字体
    # 设置加粗
    run1.font.bold=True
    # 段后距离5磅
    p1.space_after=Pt(5)
    # 段前距离5磅
    p1.space_before=Pt(5)

    p2 = document.add_paragraph()
    run2 = p2.add_run("对比结果")
    run2.font.name = '仿宋_GB2312'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run2.font.size = Pt(16)  # 设置字体
    # 设置加粗
    run2.font.bold=False

    # 添加图片   图片路径和尺寸
    document.add_picture('1.jpg', width=Inches(6))

    p3 = document.add_paragraph()
    run3 = p3.add_run("眉间与眼角之间的垂直距离/眉顶至下巴的高度=")
    run3.font.name = '仿宋_GB2312'
    run3.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run3.font.size = Pt(16)  # 设置字体
    run33 = p3.add_run("25.058%")
    red = RGBColor(255, 0, 0)
    run33.font.color.rgb = red

    p4 = document.add_paragraph()
    run4 = p4.add_run("眉间与眼角之间的垂直距离/眉顶至下巴的高度=")
    run4.font.name = '仿宋_GB2312'
    run4.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run4.font.size = Pt(16)  # 设置字体
    # 设置加粗
    run4.font.bold = False
    run44 = p4.add_run("25.058%")
    red = RGBColor(255, 0, 0)
    run44.font.color.rgb = red

    p5 = document.add_paragraph()
    run5 = p5.add_run("鼻尖与嘴角之间的垂直距离/眉顶至下巴高度=")
    run5.font.name = '仿宋_GB2312'
    run5.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run5.font.size = Pt(16)  # 设置字体
    run55 = p5.add_run("25.058%")
    red = RGBColor(255, 0, 0)
    run55.font.color.rgb = red

    p6 = document.add_paragraph()
    run6 = p6.add_run("嘴角与下巴之间的垂直距离/眉间至下巴的高度=")
    run6.font.name = '仿宋_GB2312'
    run6.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run6.font.size = Pt(16)  # 设置字体
    run66 = p6.add_run("25.058%")
    red = RGBColor(255, 0, 0)
    run66.font.color.rgb = red

    document.save('测试文档3.docx')
