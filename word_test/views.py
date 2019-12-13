from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WD_ALIGN_PARAGRAPH   # 用作设置段落对齐
from docx.shared import Pt  # 磅数
from docx.oxml.ns import qn  # 中文格式
from docx.shared import Inches  # 图片尺寸

import time
today=time.strftime("%Y-%m-%d",time.localtime())
price=100
company_list=['客户1']

if __name__ == '__main__':
	document =Document()
	document.styles['Normal'].font.name=u'微软雅黑'
	document.styles['Normal'].font.size=Pt(14)
	# 设置文档的基础字体
	document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'),u'微软雅黑')

	p1=document.add_paragraph()
	# 初始化建立第一个自然段  设置对齐方式为居中，默认为左对齐
	p1.aligment=WD_ALIGN_PARAGRAPH.CENTER
	# 标题内容
	run1=p1.add_run('在现有比对技术的基础上需要添加的人像比对后输出格式和要求')
	run1.font.name = '微软雅黑'
	run1.font.size=Pt(21) #设置字体
	# 设置加粗
	run1.font.bold=True
	# 段后距离5磅
	p1.space_after=Pt(5)
	#  #段前距离5磅
	p1.space_before=Pt(5)

	p2=document.add_paragraph()
	run2=p2.add_run("检材及样本图片的文件信息")
	run2.font.name='仿宋_GB2312'
	run2.element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋_GB2312')
	run2.font.size=Pt(16)  # 设置字体
	# 设置加粗
	run2.font.bold=False

	p3 = document.add_paragraph()
	run3 = p3.add_run("文件名称：XXXXXXXXXXXX")
	run3.font.name = '仿宋_GB2312'
	run3.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
	run3.font.size = Pt(16)  # 设置字体
	# 设置加粗
	run3.font.bold = False

	p4 = document.add_paragraph()
	run4 = p4.add_run("路径：XXXXXXXXXXXX")
	run4.font.name = '仿宋_GB2312'
	run4.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
	run4.font.size = Pt(16)  # 设置字体
	# 设置加粗
	run4.font.bold = False

	p5 = document.add_paragraph()
	run5 = p5.add_run("文件大小：XXGB(XXXX字节)")
	run5.font.name = '仿宋_GB2312'
	run5.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
	run5.font.size = Pt(16)  # 设置字体

	p6 = document.add_paragraph()
	run6 = p6.add_run("MD5哈希校验值：XXXXXXXXXXXXX")
	run6.font.name = '仿宋_GB2312'
	run6.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
	run6.font.size = Pt(16)  # 设置字体

	p7 = document.add_paragraph()
	run7 = p7.add_run("文件内容截图")
	run7.font.name = '仿宋_GB2312'
	run7.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
	run7.font.size = Pt(16)  # 设置字体
	# 添加图片   图片路径和尺寸
	document.add_picture('1.jpg', width=Inches(6))

	p8 = document.add_paragraph()
	run8 = p8.add_run("检材(比对清晰页面截图)")
	run8.font.name = '仿宋_GB2312'
	run8.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
	run8.font.size = Pt(16)  # 设置字体
	# 添加图片   图片路径和尺寸
	document.add_picture('1.jpg', width=Inches(6))

	p2 = document.add_paragraph()
	run2 = p2.add_run("局部放大锐化图")
	run2.font.name = '仿宋_GB2312'
	run2.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
	run2.font.size = Pt(16)  # 设置字体
	# 添加图片   图片路径和尺寸
	document.add_picture('1.jpg', width=Inches(6))

	p2 = document.add_paragraph()
	run2 = p2.add_run("样本图片")
	run2.font.name = '仿宋_GB2312'
	run2.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
	run2.font.size = Pt(16)  # 设置字体
	# 设置加粗
	run2.font.bold = False
	# 添加图片   图片路径和尺寸
	document.add_picture('1.jpg', width=Inches(6))


	#插入分页符
	document.add_page_break()
	document.save('测试文档.docx')#以“客户名-价格通知”作为文件名保存
