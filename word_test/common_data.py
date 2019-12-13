# coding=utf-8
# auther: wangzehua
# function: 生成word文档
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WD_ALIGN_PARAGRAPH   # 用作设置段落对齐
from docx.shared import Pt, RGBColor  # 磅数
from docx.oxml.ns import qn  # 中文格式
from docx.shared import Inches  # 图片尺寸


def r_image_create(filename, filepath, filesize, filemd5, content_img, material_img, enlarge_img, sample_img):
    """
    检材及样本图片的文件信息
    :param filename:
    :param filepath:
    :param filesize:
    :param filemd5:
    :param content_img:
    :param material_img:
    :param enlarge_img:
    :param sample_img:
    :return:
    """
    document = Document()
    document.styles['Normal'].font.name = u'微软雅黑'
    document.styles['Normal'].font.size = Pt(14)
    # 设置文档的基础字体
    document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

    p1 = document.add_paragraph()
    # 初始化建立第一个自然段  设置对齐方式为居中，默认为左对齐
    p1.aligment = WD_ALIGN_PARAGRAPH.CENTER
    # 标题内容
    run1 = p1.add_run('在现有比对技术的基础上需要添加的人像比对后输出格式和要求')
    run1.font.name = '微软雅黑'
    run1.font.size = Pt(21)  # 设置字体
    # 设置加粗
    run1.font.bold = True
    # 段后距离5磅
    p1.space_after = Pt(5)
    #  #段前距离5磅
    p1.space_before = Pt(5)

    p2 = document.add_paragraph()
    run2 = p2.add_run("检材及样本图片的文件信息")
    run2.font.name = '仿宋_GB2312'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run2.font.size = Pt(16)  # 设置字体
    # 设置加粗
    run2.font.bold = False

    p3 = document.add_paragraph()
    run3 = p3.add_run(f"文件名称：{filename}")
    run3.font.name = '仿宋_GB2312'
    run3.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run3.font.size = Pt(16)  # 设置字体
    # 设置加粗
    run3.font.bold = False

    p4 = document.add_paragraph()
    run4 = p4.add_run(f"路径：{filepath}")
    run4.font.name = '仿宋_GB2312'
    run4.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run4.font.size = Pt(16)  # 设置字体
    # 设置加粗
    run4.font.bold = False

    p5 = document.add_paragraph()
    run5 = p5.add_run(f"文件大小：{filesize}GB({filesize}字节)")
    run5.font.name = '仿宋_GB2312'
    run5.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run5.font.size = Pt(16)  # 设置字体

    p6 = document.add_paragraph()
    run6 = p6.add_run(f"MD5哈希校验值：{filemd5}")
    run6.font.name = '仿宋_GB2312'
    run6.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run6.font.size = Pt(16)  # 设置字体

    p7 = document.add_paragraph()
    run7 = p7.add_run("文件内容截图")
    run7.font.name = '仿宋_GB2312'
    run7.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run7.font.size = Pt(16)  # 设置字体
    # 添加图片   图片路径和尺寸
    document.add_picture(f"{content_img}", width=Inches(6))

    p8 = document.add_paragraph()
    run8 = p8.add_run("检材(比对清晰页面截图)")
    run8.font.name = '仿宋_GB2312'
    run8.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run8.font.size = Pt(16)  # 设置字体
    # 添加图片   图片路径和尺寸
    document.add_picture(f"{material_img}", width=Inches(6))

    p2 = document.add_paragraph()
    run2 = p2.add_run("局部放大锐化图")
    run2.font.name = '仿宋_GB2312'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run2.font.size = Pt(16)  # 设置字体
    # 添加图片   图片路径和尺寸
    document.add_picture(f"{enlarge_img}", width=Inches(6))

    p2 = document.add_paragraph()
    run2 = p2.add_run("样本图片")
    run2.font.name = '仿宋_GB2312'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run2.font.size = Pt(16)  # 设置字体
    # 设置加粗
    run2.font.bold = False
    # 添加图片   图片路径和尺寸
    document.add_picture(f"{sample_img}", width=Inches(6))

    document.save('测试文档.docx')  # 以“客户名-价格通知”作为文件名保存


def r_measure_data(file_img, brow_eye_rate, eye_nose_rate, nose_mouth_rate, mouth_chin_rate):
    """
    检材中人像头部和五官测量图
    :param file_img:
    :param brow_eye_rate:
    :param eye_nose_rate:
    :param nose_mouth_rate:
    :param mouth_chin_rate:
    :return:
    """
    document = Document()
    document.styles['Normal'].font.name = u'微软雅黑'
    document.styles['Normal'].font.size = Pt(14)
    # 设置文档的基础字体
    document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

    p1 = document.add_paragraph()
    # 初始化建立第一个自然段  设置对齐方式为居中，默认为左对齐
    p1.aligment = WD_ALIGN_PARAGRAPH.CENTER
    # 标题内容
    run1 = p1.add_run('样本中人像头部和五官测量图')
    run1.font.name = '微软雅黑'
    run1.font.size = Pt(21)  # 设置字体
    # 设置加粗
    run1.font.bold = True
    # 段后距离5磅
    p1.space_after = Pt(5)
    # 段前距离5磅
    p1.space_before = Pt(5)

    p2 = document.add_paragraph()
    run2 = p2.add_run("对比结果")
    run2.font.name = '仿宋_GB2312'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run2.font.size = Pt(16)  # 设置字体
    # 设置加粗
    run2.font.bold = False

    # 添加图片   图片路径和尺寸
    document.add_picture(f'{ file_img }', width=Inches(6))

    p3 = document.add_paragraph()
    run3 = p3.add_run("眉间与眼角之间的垂直距离/眉顶至下巴的高度=")
    run3.font.name = '仿宋_GB2312'
    run3.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run3.font.size = Pt(16)  # 设置字体
    run33 = p3.add_run(f"{brow_eye_rate}%")
    red = RGBColor(255, 0, 0)
    run33.font.color.rgb = red

    p4 = document.add_paragraph()
    run4 = p4.add_run("眉间与眼角之间的垂直距离/眉顶至下巴的高度=")
    run4.font.name = '仿宋_GB2312'
    run4.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run4.font.size = Pt(16)  # 设置字体
    # 设置加粗
    run4.font.bold = False
    run44 = p4.add_run(f"{eye_nose_rate}%")
    red = RGBColor(255, 0, 0)
    run44.font.color.rgb = red

    p5 = document.add_paragraph()
    run5 = p5.add_run("鼻尖与嘴角之间的垂直距离/眉顶至下巴高度=")
    run5.font.name = '仿宋_GB2312'
    run5.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run5.font.size = Pt(16)  # 设置字体
    run55 = p5.add_run(f"{nose_mouth_rate}%")
    red = RGBColor(255, 0, 0)
    run55.font.color.rgb = red

    p6 = document.add_paragraph()
    run6 = p6.add_run("嘴角与下巴之间的垂直距离/眉间至下巴的高度=")
    run6.font.name = '仿宋_GB2312'
    run6.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run6.font.size = Pt(16)  # 设置字体
    run66 = p6.add_run(f"{mouth_chin_rate}%")
    red = RGBColor(255, 0, 0)
    run66.font.color.rgb = red

    document.save('测试文档3.docx')


def r_table_data():
    document = Document()
    document.styles['Normal'].font.name = u'微软雅黑'
    document.styles['Normal'].font.size = Pt(14)
    # 设置文档的基础字体
    document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

    p1 = document.add_paragraph()
    # 初始化建立第一个自然段  设置对齐方式为居中，默认为左对齐
    p1.aligment = WD_ALIGN_PARAGRAPH.CENTER
    # 标题内容
    run1 = p1.add_run('人像细节特征比较图：(主要五官的特征比对)')
    run1.font.name = '微软雅黑'
    run1.font.size = Pt(21)  # 设置字体
    # 设置加粗
    run1.font.bold = True
    # 段后距离5磅
    p1.space_after = Pt(5)
    ##段前距离5磅
    p1.space_before = Pt(5)

    table = document.add_table(rows=5, cols=4, style='Table Grid')

    table.cell(0, 0).text = '名称'
    table.cell(0, 1).text = '检材细节特征'
    table.cell(0, 2).text = '样本细节特征'
    table.cell(0, 3).text = '细节特征一致性对比结果'
    table.cell(1, 0).text = '眼睛及眉毛特征'
    table.cell(2, 0).text = '鼻部及人中特征'
    table.cell(3, 0).text = '唇部特征'
    table.cell(4, 0).text = '耳部特征'
    run = document.tables[0].cell(1, 2).paragraphs[0].add_run()
    run.add_picture('2.png')

    document.save('测试文档2.docx')  # 以“客户名-价格通知”作为文件名保存